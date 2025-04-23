import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from io import BytesIO
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

st.title("📝 Generador de Oficios")

# 1️⃣ Número de oficio
numero_oficio = st.text_input("Número de oficio", placeholder="Ej. OF-123/2025")

# 2️⃣ Fecha con calendario
fecha_seleccionada = st.date_input("Selecciona la fecha", value=datetime.date.today())
fecha = fecha_seleccionada.strftime("%d DE %B DE %Y").upper()

# 3️⃣ Asunto
asuntos = [
    "SOLICITUD DE INFORMACIÓN",
    "SEGUIMIENTO A PROYECTO",
    "ENTREGA DE RESULTADOS",
    "CITATORIO A REUNIÓN"
]
asunto = st.selectbox("Selecciona el asunto", asuntos)

# 4️⃣ Destinatario
destinatarios = [
    {"nombre": "Juan Pérez", "cargo": "Director General"},
    {"nombre": "Laura Gómez", "cargo": "Jefa de Finanzas"},
    {"nombre": "Carlos Ruiz", "cargo": "Coordinador de Proyectos"}
]

nombres = [d["nombre"] for d in destinatarios]
nombre_seleccionado = st.selectbox("Selecciona un destinatario", nombres)
destinatario = next(d for d in destinatarios if d["nombre"] == nombre_seleccionado)

# 5️⃣ Inputs dinámicos
if "textos" not in st.session_state:
    st.session_state.textos = []
if "tablas" not in st.session_state:
    st.session_state.tablas = []

if st.button("Agregar texto"):
    st.session_state.textos.append("")

if st.button("Agregar tabla"):
    st.session_state.tablas.append([["", ""]])

# Campos de texto
for i, texto in enumerate(st.session_state.textos):
    st.session_state.textos[i] = st.text_area(f"Texto #{i+1}", value=texto)

# Tablas
for i, tabla in enumerate(st.session_state.tablas):
    st.markdown(f"**Tabla #{i+1}**")
    cols = st.columns(2)
    fila1 = cols[0].text_input(f"T{i}_0", value=tabla[0][0], key=f"t{i}_00")
    fila2 = cols[1].text_input(f"T{i}_1", value=tabla[0][1], key=f"t{i}_01")
    st.session_state.tablas[i] = [[fila1, fila2]]

# 6️⃣ Firma
firmante_nombre = st.text_input("Nombre del firmante", placeholder="Ej. María López")
firmante_cargo = st.text_input("Cargo del firmante", placeholder="Ej. Subdirectora de Administración")

# 7️⃣ CC
cc = st.text_area("CC (Con Copia Para)", placeholder="Ej. Archivo, Coordinación de Comunicación, etc.")

# 8️⃣ Generar documento
if st.button("Generar oficio"):
    doc = Document("plantilla_file.docx")

    section = doc.sections[0]
    header = section.header

    # Encabezado institucional + fecha + asunto
    parrafo_derecha = header.add_paragraph()
    parrafo_derecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = parrafo_derecha.add_run(f"DIRECCIÓN DE ADMINISTRACIÓN Y FINANZAS\nCIUDAD DE MÉXICO, {fecha}\nASUNTO: {asunto}")
    run.bold = True
    run.font.size = Pt(11)

    # Número de oficio
    num_parrafo = header.add_paragraph()
    num_parrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = num_parrafo.add_run(numero_oficio)
    run.bold = True
    run.font.size = Pt(12)

    # Cuerpo del oficio
    doc.add_paragraph(f"Destinatario: {destinatario['nombre']}, {destinatario['cargo']}")

    for texto in st.session_state.textos:
        doc.add_paragraph(texto)

    for tabla in st.session_state.tablas:
        table = doc.add_table(rows=1, cols=len(tabla[0]))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(tabla[0]):
            hdr_cells[i].text = col

    # Firma
    doc.add_paragraph("\n\nAtentamente,\n")
    doc.add_paragraph(firmante_nombre)
    doc.add_paragraph(firmante_cargo)

    # CC
    if cc.strip():
        doc.add_paragraph("\nC.C.P. " + cc, style='Normal')

    # Guardar y descargar
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("Oficio generado correctamente 🎉")
    st.download_button("📥 Descargar oficio", buffer, file_name=f"Oficio_{numero_oficio}.docx")
