from docx import Document
from io import BytesIO
import zipfile

def generar_oficios(df, plantilla_file):
    zip_buffer = BytesIO()
    plantilla = Document(plantilla_file)

    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zipf:
        for _, row in df.iterrows():
            doc = Document(plantilla_file)  # Usar una copia para cada uno

            for p in doc.paragraphs:
                for key in df.columns:
                    placeholder = f"<<{key}>>"
                    if placeholder in p.text:
                        p.text = p.text.replace(placeholder, str(row[key]))

            buffer = BytesIO()
            doc.save(buffer)
            zipf.writestr(f"Oficio_{row['Nombre']}.docx", buffer.getvalue())

    zip_buffer.seek(0)
    return zip_buffer
